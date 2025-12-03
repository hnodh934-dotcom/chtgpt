#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
نظام توليد النماذج الاحترافية - شركة وثقى للاستشارات المهنية
Wathqa Professional Document Generator v2.0
"""

import os
import uuid
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════
#                         إعدادات النظام
# ═══════════════════════════════════════════════════════════════════

OUTPUT_DIR = "/home/user/chtgpt/wathqa_output/"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ═══════════════════════════════════════════════════════════════════
#                     الهوية البصرية الفاخرة
# ═══════════════════════════════════════════════════════════════════

# الألوان الأساسية
COLOR_PRIMARY = RGBColor(27, 94, 79)      # #1B5E4F - أخضر داكن فاخر
COLOR_GOLD = RGBColor(197, 164, 103)      # #C5A467 - ذهبي كلاسيكي
COLOR_GOLD_DARK = RGBColor(166, 134, 71)  # #A68647 - ذهبي داكن
COLOR_GOLD_LIGHT = RGBColor(218, 194, 148) # #DAC294 - ذهبي فاتح
COLOR_BLACK = RGBColor(33, 33, 33)        # #212121 - أسود ناعم
COLOR_GRAY = RGBColor(117, 117, 117)      # #757575 - رمادي
COLOR_LIGHT_BG = RGBColor(250, 248, 245)  # #FAF8F5 - خلفية كريمية
COLOR_WHITE = RGBColor(255, 255, 255)     # #FFFFFF - أبيض

# الخطوط
FONT_MAIN = "Sakkal Majalla"
FONT_ENGLISH = "Georgia"

# ═══════════════════════════════════════════════════════════════════
#                     نظام الترقيم الفريد
# ═══════════════════════════════════════════════════════════════════

class CaseNumberGenerator:
    """
    مولّد أرقام القضايا الفريدة
    الصيغة: WQ-YYYY-MMDD-XXXX
    """

    @staticmethod
    def generate(prefix="WQ"):
        now = datetime.now()
        date_part = now.strftime("%Y-%m%d")

        # توليد رقم فريد من 4 خانات
        unique_hash = hashlib.md5(
            f"{now.timestamp()}{uuid.uuid4()}".encode()
        ).hexdigest()[:4].upper()

        return f"{prefix}-{date_part}-{unique_hash}"

    @staticmethod
    def generate_full():
        """رقم كامل مع الوقت"""
        now = datetime.now()
        return f"WQ-{now.strftime('%Y%m%d')}-{now.strftime('%H%M%S')}-{uuid.uuid4().hex[:4].upper()}"

# ═══════════════════════════════════════════════════════════════════
#                     الدوال المساعدة الأساسية
# ═══════════════════════════════════════════════════════════════════

def set_cell_shading(cell, hex_color):
    """تلوين خلفية الخلية"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C5A467", size="6"):
    """إضافة حدود ذهبية فاخرة للخلية"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)

    tcPr.append(tcBorders)

def set_cell_rtl(cell):
    """ضبط اتجاه الخلية من اليمين لليسار"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # إضافة خاصية RTL للفقرة
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

def set_paragraph_rtl(paragraph):
    """ضبط اتجاه الفقرة من اليمين لليسار"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def configure_document(doc):
    """إعداد المستند الأساسي"""
    # إعداد الستايل الافتراضي
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_MAIN
    font.size = Pt(13)

    # إعداد الخط العربي
    style.element.rPr.rFonts.set(qn('w:ascii'), FONT_MAIN)
    style.element.rPr.rFonts.set(qn('w:hAnsi'), FONT_MAIN)
    style.element.rPr.rFonts.set(qn('w:cs'), FONT_MAIN)

    # إعداد اتجاه المستند RTL
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

    return doc

# ═══════════════════════════════════════════════════════════════════
#                     عناصر التصميم الفاخر
# ═══════════════════════════════════════════════════════════════════

def add_golden_line(doc, width=6):
    """إضافة خط ذهبي فاصل"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # إنشاء خط ذهبي باستخدام الحدود
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width * 4))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C5A467')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return p

def add_double_golden_line(doc):
    """إضافة خط ذهبي مزدوج"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    # الخط العلوي
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'C5A467')
    pBdr.append(top)

    # الخط السفلي
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B5E4F')
    pBdr.append(bottom)

    pPr.append(pBdr)
    return p

def create_luxury_header(doc, title, form_code, case_number):
    """
    إنشاء ترويسة فاخرة مع الخط الذهبي
    """
    # ═══ اسم الشركة ═══
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = company.add_run("━━━━━━━━━━  ")
    run.font.color.rgb = COLOR_GOLD
    run.font.size = Pt(14)

    run2 = company.add_run("شركة وثقى للاستشارات المهنية")
    run2.bold = True
    run2.font.size = Pt(22)
    run2.font.color.rgb = COLOR_PRIMARY

    run3 = company.add_run("  ━━━━━━━━━━")
    run3.font.color.rgb = COLOR_GOLD
    run3.font.size = Pt(14)

    # ═══ الشعار الفرعي ═══
    slogan = doc.add_paragraph()
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_slogan = slogan.add_run("WATHQA Professional Consultancy")
    run_slogan.font.size = Pt(10)
    run_slogan.font.color.rgb = COLOR_GRAY
    run_slogan.font.name = FONT_ENGLISH
    run_slogan.italic = True

    # ═══ الخط الذهبي الفاخر ═══
    add_golden_line(doc, width=3)

    # ═══ عنوان النموذج ═══
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)

    # رمز زخرفي
    deco1 = title_p.add_run("◆ ")
    deco1.font.color.rgb = COLOR_GOLD
    deco1.font.size = Pt(16)

    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = COLOR_BLACK

    deco2 = title_p.add_run(" ◆")
    deco2.font.color.rgb = COLOR_GOLD
    deco2.font.size = Pt(16)

    # ═══ صندوق معلومات النموذج ═══
    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # رقم النموذج (يسار)
    cell_left = info_table.cell(0, 0)
    cell_left.text = f"رقم النموذج: {form_code}"
    for p in cell_left.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_GRAY

    # التاريخ (وسط)
    cell_center = info_table.cell(0, 1)
    today = datetime.now().strftime("%Y/%m/%d")
    cell_center.text = f"التاريخ: {today}"
    for p in cell_center.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_GRAY

    # رقم القضية الفريد (يمين)
    cell_right = info_table.cell(0, 2)
    cell_right.text = f"الرقم المرجعي: {case_number}"
    for p in cell_right.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_GOLD_DARK
            r.bold = True

    # خط ذهبي مزدوج
    add_double_golden_line(doc)
    doc.add_paragraph()  # مسافة

def add_section_title(doc, title, number=None):
    """إضافة عنوان قسم فاخر"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)

    # الرقم
    if number:
        num_run = p.add_run(f"{number}. ")
        num_run.font.color.rgb = COLOR_GOLD_DARK
        num_run.font.size = Pt(14)
        num_run.bold = True

    # العنوان
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = COLOR_PRIMARY

    # خط تحت العنوان
    underline = p.add_run("  ─────────")
    underline.font.color.rgb = COLOR_GOLD_LIGHT
    underline.font.size = Pt(10)

    set_paragraph_rtl(p)
    return p

def create_luxury_table(doc, rows, cols, headers=None):
    """إنشاء جدول فاخر مع تنسيق ذهبي"""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # اتجاه الجدول RTL
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)

    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tblPr.append(bidi)

    # تنسيق الخلايا
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # الحدود الذهبية
            set_cell_borders(cell, color="C5A467", size="4")

            # المحاذاة العمودية
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # اتجاه النص RTL
            set_cell_rtl(cell)

            # تنسيق الصف الأول (الهيدر)
            if i == 0 and headers:
                set_cell_shading(cell, "1B5E4F")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.color.rgb = COLOR_WHITE
                        r.bold = True
                        r.font.size = Pt(12)

            # تناوب ألوان الصفوف
            elif i > 0:
                if i % 2 == 0:
                    set_cell_shading(cell, "FAF8F5")
                else:
                    set_cell_shading(cell, "FFFFFF")

    # تعبئة الهيدر
    if headers:
        for idx, header in enumerate(headers):
            if idx < cols:
                table.cell(0, idx).text = header
                set_cell_rtl(table.cell(0, idx))

    return table

def add_luxury_signature_section(doc, signatures):
    """قسم التوقيعات الفاخر"""
    doc.add_paragraph()
    add_golden_line(doc, width=2)

    table = doc.add_table(rows=3, cols=len(signatures))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, title in enumerate(signatures):
        # خط التوقيع
        cell_line = table.cell(0, i)
        p_line = cell_line.paragraphs[0]
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line = p_line.add_run("▬▬▬▬▬▬▬▬▬▬▬▬")
        run_line.font.color.rgb = COLOR_GOLD
        run_line.font.size = Pt(10)

        # العنوان
        cell_title = table.cell(1, i)
        p_title = cell_title.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = COLOR_PRIMARY

        # مكان الكتابة
        cell_space = table.cell(2, i)
        p_space = cell_space.paragraphs[0]
        p_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_space = p_space.add_run("........................")
        run_space.font.color.rgb = COLOR_GRAY
        run_space.font.size = Pt(10)

def add_footer_branding(doc):
    """إضافة ذيل فاخر للمستند"""
    doc.add_paragraph()
    doc.add_paragraph()
    add_golden_line(doc, width=2)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = footer.add_run("شركة وثقى للاستشارات المهنية")
    run1.font.size = Pt(9)
    run1.font.color.rgb = COLOR_GRAY

    run2 = footer.add_run("  ◆  ")
    run2.font.color.rgb = COLOR_GOLD
    run2.font.size = Pt(8)

    run3 = footer.add_run("جميع الحقوق محفوظة")
    run3.font.size = Pt(9)
    run3.font.color.rgb = COLOR_GRAY

# ═══════════════════════════════════════════════════════════════════
#                     النماذج الاحترافية
# ═══════════════════════════════════════════════════════════════════

def create_wq001():
    """نموذج استلام وثائق ومستندات"""
    doc = Document()
    configure_document(doc)
    case_num = CaseNumberGenerator.generate()

    create_luxury_header(doc, "نموذج استلام وثائق ومستندات", "WQ-001", case_num)

    # ═══ 1. بيانات العميل والطلب ═══
    add_section_title(doc, "بيانات العميل والطلب", "1")

    table1 = create_luxury_table(doc, rows=4, cols=4)

    # تعبئة البيانات (RTL - من اليمين)
    data = [
        ("التاريخ:", "", "وقت الاستلام:", ""),
        ("اسم العميل:", "", "رقم الجوال:", ""),
        ("نوع الخدمة:", "", "البريد الإلكتروني:", ""),
        ("المرجع:", "", "رقم الملف:", "")
    ]

    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, 3-j)  # عكس لـ RTL
            cell.text = val
            set_cell_rtl(cell)
            # تلوين التسميات
            if j % 2 == 0:
                set_cell_shading(cell, "FAF8F5")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    # ═══ 2. قائمة التحقق ═══
    add_section_title(doc, "قائمة التحقق من المستندات", "2")

    note = doc.add_paragraph()
    note_run = note.add_run("◀ يرجى التأشير بعلامة (✓) أمام المستند المستلم")
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = COLOR_GRAY
    note_run.italic = True
    set_paragraph_rtl(note)

    headers = ["ملاحظات", "الحالة", "المستند المطلوب", "م"]
    table2 = create_luxury_table(doc, rows=1, cols=4, headers=headers)

    items = [
        "الهوية الوطنية / الإقامة / الجواز",
        "السجل التجاري / عقد التأسيس",
        "الصكوك / عقود الإيجار / إثبات الملكية",
        "المستندات الأساسية للدعوى",
        "الوكالات السابقة / حصر الورثة",
        "المراسلات والمخاطبات السابقة",
        "مستندات أخرى"
    ]

    for idx, item in enumerate(items, 1):
        row = table2.add_row()
        row.cells[3].text = str(idx)
        row.cells[2].text = item
        row.cells[1].text = "☐"
        row.cells[0].text = ""

        for cell in row.cells:
            set_cell_borders(cell, color="C5A467", size="4")
            set_cell_rtl(cell)
            if idx % 2 == 0:
                set_cell_shading(cell, "FAF8F5")

    doc.add_paragraph()

    # ═══ 3. إقرار الاستلام ═══
    add_section_title(doc, "إقرار الاستلام", "3")

    declaration = doc.add_paragraph()
    dec_run = declaration.add_run(
        "أقر أنا الموظف المختص باستلام المستندات الموضحة أعلاه، "
        "وتم إيداعها في ملف العميل الرقمي/الورقي."
    )
    dec_run.font.size = Pt(11)
    set_paragraph_rtl(declaration)

    add_luxury_signature_section(doc, ["اسم المستلم", "التوقيع", "التاريخ"])
    add_footer_branding(doc)

    doc.save(f"{OUTPUT_DIR}WQ-001_استلام_وثائق_{case_num}.docx")
    print(f"✅ تم إنشاء: WQ-001 - الرقم: {case_num}")
    return case_num


def create_wq002():
    """تقرير فحص الوقائع والأوراق"""
    doc = Document()
    configure_document(doc)
    case_num = CaseNumberGenerator.generate()

    create_luxury_header(doc, "تقرير فحص الوقائع والأوراق", "WQ-002", case_num)

    # ═══ 1. ملخص ملف القضية ═══
    add_section_title(doc, "ملخص ملف القضية", "1")

    table1 = create_luxury_table(doc, rows=3, cols=4)

    labels = [
        ("اسم العميل:", "", "رقم الملف:", ""),
        ("الموضوع:", "", "تاريخ الفحص:", ""),
        ("المستشار المسؤول:", "", "درجة الأهمية:", "")
    ]

    for i, row_data in enumerate(labels):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, 3-j)
            cell.text = val
            set_cell_rtl(cell)
            if j % 2 == 0:
                set_cell_shading(cell, "FAF8F5")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    # ═══ 2. السرد الوقائي ═══
    add_section_title(doc, "السرد الوقائي والتحليل", "2")

    for i in range(1, 5):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ...................................................................")
        run.font.color.rgb = COLOR_GRAY
        set_paragraph_rtl(p)

    doc.add_paragraph()

    # ═══ 3. فحص الأسانيد ═══
    add_section_title(doc, "فحص الأسانيد والمستندات", "3")

    headers = ["ملاحظات المستشار", "الوزن القانوني", "الحالة", "المستند", "م"]
    table3 = create_luxury_table(doc, rows=1, cols=5, headers=headers)

    for i in range(1, 5):
        row = table3.add_row()
        row.cells[4].text = str(i)
        for cell in row.cells:
            set_cell_borders(cell, color="C5A467", size="4")
            set_cell_rtl(cell)
            if i % 2 == 0:
                set_cell_shading(cell, "FAF8F5")

    doc.add_paragraph()

    # ═══ 4. الرأي الفني ═══
    add_section_title(doc, "الرأي الفني الأولي", "4")

    p1 = doc.add_paragraph()
    run1 = p1.add_run("◀ التوصية: ")
    run1.bold = True
    run1.font.color.rgb = COLOR_PRIMARY
    p1.add_run("........................................................................")
    set_paragraph_rtl(p1)

    p2 = doc.add_paragraph()
    run2 = p2.add_run("◀ النواقص المطلوبة: ")
    run2.bold = True
    run2.font.color.rgb = COLOR_PRIMARY
    p2.add_run("...................................................................")
    set_paragraph_rtl(p2)

    add_luxury_signature_section(doc, ["توقيع المستشار", "الاعتماد"])
    add_footer_branding(doc)

    doc.save(f"{OUTPUT_DIR}WQ-002_فحص_الوقائع_{case_num}.docx")
    print(f"✅ تم إنشاء: WQ-002 - الرقم: {case_num}")
    return case_num


def create_wq003():
    """دراسة المسارات الاستراتيجية"""
    doc = Document()
    configure_document(doc)
    case_num = CaseNumberGenerator.generate()

    create_luxury_header(doc, "دراسة المسارات الاستراتيجية", "WQ-003", case_num)

    # ═══ 1. بطاقة المشروع ═══
    add_section_title(doc, "بطاقة المشروع", "1")

    table1 = create_luxury_table(doc, rows=2, cols=4)

    data = [
        ("اسم العميل:", "", "رقم الملف:", ""),
        ("نوع القضية:", "", "تاريخ العرض:", "")
    ]

    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, 3-j)
            cell.text = val
            set_cell_rtl(cell)
            if j % 2 == 0:
                set_cell_shading(cell, "FAF8F5")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = COLOR_PRIMARY

    p_desc = doc.add_paragraph()
    run_desc = p_desc.add_run("◀ وصف الحالة: ")
    run_desc.bold = True
    run_desc.font.color.rgb = COLOR_PRIMARY
    p_desc.add_run("...................................................................")
    set_paragraph_rtl(p_desc)

    doc.add_paragraph()

    # ═══ 2. المسارات القانونية ═══
    add_section_title(doc, "المسارات القانونية المقترحة", "2")

    headers = ["المسار البديل", "المسار الموصى به", "عنصر المقارنة"]
    table2 = create_luxury_table(doc, rows=1, cols=3, headers=headers)

    criteria = ["الإجراء", "المدة المتوقعة", "التكلفة التقديرية", "نسبة النجاح"]

    for idx, c in enumerate(criteria, 1):
        row = table2.add_row()
        row.cells[2].text = c
        for cell in row.cells:
            set_cell_borders(cell, color="C5A467", size="4")
            set_cell_rtl(cell)
            if idx % 2 == 0:
                set_cell_shading(cell, "FAF8F5")
        # تمييز عمود المعايير
        set_cell_shading(row.cells[2], "FAF8F5")
        for p in row.cells[2].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # ═══ 3. تحليل المخاطر ═══
    add_section_title(doc, "تحليل المخاطر والمكتسبات", "3")

    table3 = create_luxury_table(doc, rows=1, cols=2, headers=["المخاطر المحتملة", "نقاط القوة"])

    for i in range(3):
        row = table3.add_row()
        for cell in row.cells:
            set_cell_borders(cell, color="C5A467", size="4")
            set_cell_rtl(cell)

    doc.add_paragraph()

    # ═══ 4. التوصية النهائية ═══
    add_section_title(doc, "التوصية النهائية", "4")

    rec = doc.add_paragraph()
    rec_run = rec.add_run(
        "بناءً على الدراسة أعلاه، توصي شركة وثقى باتباع المسار (...............) للأسباب التالية:"
    )
    set_paragraph_rtl(rec)

    for i in range(1, 4):
        p = doc.add_paragraph()
        p.add_run(f"    {i}. ................................................................")
        set_paragraph_rtl(p)

    add_luxury_signature_section(doc, ["إعداد المستشار", "التاريخ"])
    add_footer_branding(doc)

    doc.save(f"{OUTPUT_DIR}WQ-003_المسارات_الاستراتيجية_{case_num}.docx")
    print(f"✅ تم إنشاء: WQ-003 - الرقم: {case_num}")
    return case_num


def create_wq004():
    """طلب إصدار وكالة شرعية"""
    doc = Document()
    configure_document(doc)
    case_num = CaseNumberGenerator.generate()

    create_luxury_header(doc, "طلب إصدار وكالة شرعية", "WQ-004", case_num)

    # ═══ ملاحظة تمهيدية ═══
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note_run = note_p.add_run(
        "◆ عزيزي العميل، لتمكيننا من تمثيلكم، نأمل إصدار وكالة عبر بوابة ناجز ◆"
    )
    note_run.font.size = Pt(11)
    note_run.font.color.rgb = COLOR_GOLD_DARK
    note_run.italic = True

    doc.add_paragraph()

    # ═══ 1. بيانات الوكيل ═══
    add_section_title(doc, "بيانات الوكيل", "1")

    table1 = create_luxury_table(doc, rows=3, cols=2)

    data = [
        ("شركة وثقى للاستشارات المهنية", "اسم المنشأة / الوكيل:"),
        ("☐ منشآت    ☐ أفراد", "نوع الوكالة:"),
        ("", "رقم الترخيص/الهوية:")
    ]

    for i, (val, label) in enumerate(data):
        table1.cell(i, 1).text = label
        table1.cell(i, 0).text = val

        set_cell_shading(table1.cell(i, 1), "FAF8F5")
        set_cell_rtl(table1.cell(i, 1))
        set_cell_rtl(table1.cell(i, 0))

        for p in table1.cell(i, 1).paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    # ═══ 2. بنود الوكالة ═══
    add_section_title(doc, "بنود الوكالة المطلوبة", "2")

    sub_note = doc.add_paragraph()
    sub_run = sub_note.add_run("◀ يرجى التأكد من تفعيل الصلاحيات التالية:")
    sub_run.font.color.rgb = COLOR_GRAY
    sub_run.italic = True
    set_paragraph_rtl(sub_note)

    powers = [
        ("أ", "صلاحيات المرافعة", [
            "☐ المرافعة عن المنشأة/الفرد في الدعاوى",
            "☐ تقديم المذكرات واللوائح والاعتراضات",
            "☐ سماع الدعاوى والرد عليها",
            "☐ المرافعة أمام جميع درجات التقاضي"
        ]),
        ("ب", "صلاحيات الإجراءات", [
            "☐ استلام وتسليم المعاملات والصكوك",
            "☐ مراجعة أقسام الشرطة والنيابة",
            "☐ مراجعة الخبراء وهيئات النظر"
        ]),
        ("ج", "صلاحيات خاصة", [
            "☐ الإقرار / الإنكار",
            "☐ الصلح / التنازل",
            "☐ قبض المبالغ / الشيكات"
        ])
    ]

    for letter, title, items in powers:
        # عنوان القسم الفرعي
        p_title = doc.add_paragraph()

        letter_run = p_title.add_run(f"{letter}) ")
        letter_run.font.color.rgb = COLOR_GOLD_DARK
        letter_run.bold = True

        title_run = p_title.add_run(title)
        title_run.bold = True
        title_run.font.color.rgb = COLOR_PRIMARY

        set_paragraph_rtl(p_title)

        # البنود
        for item in items:
            p_item = doc.add_paragraph()
            p_item.paragraph_format.left_indent = Cm(1)
            item_run = p_item.add_run(f"    {item}")
            item_run.font.size = Pt(11)
            set_paragraph_rtl(p_item)

    doc.add_paragraph()

    # ═══ 3. معلومات إضافية ═══
    add_section_title(doc, "معلومات إضافية", "3")

    info_items = [
        "◀ مدة الوكالة المقترحة: ☐ سنة   ☐ سنتين   ☐ 5 سنوات",
        "◀ رقم ملف العميل لدينا: .........................................."
    ]

    for item in info_items:
        p = doc.add_paragraph()
        p.add_run(item)
        set_paragraph_rtl(p)

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("◆ لأي استفسار، يرجى التواصل مع منسق العملاء ◆")
    contact_run.font.color.rgb = COLOR_GRAY
    contact_run.italic = True
    contact_run.font.size = Pt(10)

    add_footer_branding(doc)

    doc.save(f"{OUTPUT_DIR}WQ-004_طلب_وكالة_{case_num}.docx")
    print(f"✅ تم إنشاء: WQ-004 - الرقم: {case_num}")
    return case_num


# ═══════════════════════════════════════════════════════════════════
#                         التنفيذ الرئيسي
# ═══════════════════════════════════════════════════════════════════

def main():
    """تشغيل مولّد النماذج"""
    print("\n" + "═" * 60)
    print("     نظام توليد النماذج - شركة وثقى للاستشارات المهنية")
    print("═" * 60 + "\n")

    print(f"📁 مجلد الإخراج: {OUTPUT_DIR}\n")

    # توليد جميع النماذج
    nums = []
    nums.append(("WQ-001", create_wq001()))
    nums.append(("WQ-002", create_wq002()))
    nums.append(("WQ-003", create_wq003()))
    nums.append(("WQ-004", create_wq004()))

    print("\n" + "─" * 60)
    print("📋 ملخص الأرقام المرجعية المُنشأة:")
    print("─" * 60)
    for form, num in nums:
        print(f"   {form}: {num}")

    print("\n" + "═" * 60)
    print("     ✨ تم إنشاء جميع النماذج بنجاح ✨")
    print("═" * 60 + "\n")

if __name__ == "__main__":
    main()
